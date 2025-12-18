import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime  # ← ADD THIS LINE

load_dotenv()

API_KEY = os.getenv("OPENAI_API_KEY")

if not API_KEY:
    raise RuntimeError("OPENAI_API_KEY is missing from .env file")

client = OpenAI(api_key=API_KEY)


def generate_activity_report(activities, department_name, start_date, end_date):
    """Generate formatted report using OpenAI (v2 compatible)"""

    grouped = {
        'Departmental Activities': [],
        'Training and Placement': [],
        'Student Activities': [],
        'Student Achievements': [],
        'Other Activities': []
    }

    for activity in activities:
        activity_type = activity['Activity_Type']
        if activity_type in grouped:
            grouped[activity_type].append(activity)

    prompt = f"""
Generate a professional departmental activity report for {department_name}.
Period: {start_date} to {end_date}

Format each activity as a clear, grammatically correct sentence.
Include: activity title, description, date, organizer, and participant count.

Activities data:
{grouped}

Organize them under these headings:
1. Departmental Activities
2. Training and Placement Activities
3. Students' Activities
4. Students' Achievements
5. Other Activities

Use professional language and proper formatting.
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "You are a professional academic report writer."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return response.choices[0].message.content


def create_word_document(report_text, department_name, filename):
    """Create Word document from report text"""

    doc = Document()

    title = doc.add_heading(f'{department_name} - Activity Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in report_text.split("\n"):
        if line.strip():
            if line.startswith("#"):
                doc.add_heading(line.replace("#", "").strip(), level=1)
            else:
                doc.add_paragraph(line)

    filepath = f"uploads/reports/{filename}"
    os.makedirs("uploads/reports", exist_ok=True)
    doc.save(filepath)

    return filepath


def generate_report_from_sections(department_id, department_name, start_date, end_date):
    """
    Generate report from uploaded section files
    
    Reads files uploaded in each section and creates a summary report
    """
    
    from utils.database import execute_query
    
    query = """
        SELECT 
            s.Activity_Section,
            GROUP_CONCAT(f.File_Name SEPARATOR ', ') as files,
            COUNT(f.File_ID) as file_count,
            MAX(f.Upload_Date) as latest_upload
        FROM ACTIVITY_SECTION s
        LEFT JOIN ACTIVITY_FILE f ON s.Section_ID = f.Section_ID
        WHERE s.Department_ID = %s
        AND (f.Upload_Date IS NULL OR f.Upload_Date BETWEEN %s AND %s)
        GROUP BY s.Activity_Section
        ORDER BY FIELD(s.Activity_Section,
            'Departmental Activities',
            'Training and Placement',
            'Student Activities',
            'Student Achievements',
            'Other Activities')
    """
    
    sections_data = execute_query(query, (department_id, start_date, end_date))
    
    report = f"""
{department_name} - Activity Report
Period: {start_date} to {end_date}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

"""
    
    if sections_data:
        for section_data in sections_data:
            report += f"\n{section_data['Activity_Section']}:\n"
            if section_data['file_count'] and section_data['file_count'] > 0:
                report += f"  Files uploaded: {section_data['files']}\n"
                report += f"  Total files: {section_data['file_count']}\n"
                report += f"  Last updated: {section_data['latest_upload']}\n"
            else:
                report += f"  No files uploaded\n"
    else:
        report += "\nNo activity data found for the selected period.\n"
    
    return report